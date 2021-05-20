import csv
from os import remove
try:
    import docx
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Para Usar Este Módulo has de instalar docx con el comando pip install python-docx")
try:
    import xlsxwriter
except ImportError:
    print("Para Usar Este Módulo has de instalar el módulo xlsxwriter con el comando pip install xlsxwriter")
try:
    import xlrd
except ImportError:
    print("Para Usar Este Módulo has de instalar el módulo xlrd con el comando pip install xlrd")
from tkinter.filedialog import askopenfilename, asksaveasfilename
try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    print("Para Usar Este Módulo has de instalar el módulo PIL con el comando pip install pillow")
try:
    from docx2pdf import convert
except ImportError:
    print("Para Usar Este Módulo has de instalar docx2pdf con el comando pip install docx2pdf")
from json import dumps
try:
    from xmltodict import parse as xmlparse
except ImportError:
    print("Para Usar Este Módulo has de instalar xmltodict con el comando pip install xmltodict")
try:
    from json2xml import json2xml
    from json2xml.utils import readfromjson
except ImportError:
    print("Para Usar Este Módulo has de instalar json2xml con el comando pip install json2xml")


# Class
class HTML:
    def __init__(self, nombre, texto, descripcion="", lenguaje="es"):
        self.nombre = nombre
        self.texto = texto
        self.lenguaje = lenguaje
        self.descripcion = descripcion

    def nuevo_html(self):
        html = open("{}.html".format(str(self.nombre)), 'a')
        html.write('<!DOCTYPE html><html lang="{}"><head><meta charset="UTF-8"><meta name="description" content="{}"><title>{}</title></head><body><p>{}</p></body></html>'.format(str(self.lenguaje), str(self.descripcion), str(self.nombre), str(self.texto)).replace("%NuevaLinea%", "\n"))
        html.close()


class PNG:
    def __init__(self, name, width, height):
        self.name = name
        self.width = width
        self.height = height

    def new_one_color(self, color):
        img = Image.new('RGB', (int(self.width), int(self.height)), color=str(color))
        img.save('{}.png'.format(str(self.name)))

    def new_one_color_text(self, color, texto, color_texto):
        img = Image.new('RGB', (100, 30), color=(color[0], color[1], color[2]))
        d = ImageDraw.Draw(img)
        d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        img.save('{}.png'.format(str(self.name)))

    def new_one_color_text_font(self, color, texto, color_texto, tamano_de_texto, tu_archivo_de_fuente):
        img = Image.new('RGB', (300, 100), color=(color[0], color[1], color[2]))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        d.text((10, 10), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        img.save('{}.png'.format(str(self.name)))

    def new_one_color_insert_img(self, color, position):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.png'.format(str(self.name)))

    def new_one_color_insert_img_sin_input(self, color, position, inserted):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.png'.format(str(self.name)))

    def new_one_color_insert_img_with_text(self, color, position, texto, color_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        d = ImageDraw.Draw(img)
        try:
            d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((10, 10), str(texto), fill=(str(color_texto)))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.png'.format(str(self.name)))

    def new_one_color_insert_img_sin_input_with_text(self, color, position, inserted, texto, color_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        d = ImageDraw.Draw(img)
        try:
            d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((10, 10), str(texto), fill=(str(color_texto)))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.png'.format(str(self.name)))

    def new_one_color_insert_img_with_text_font(self, color, position, texto, position_texto, color_texto, tu_archivo_de_fuente, tamano_de_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        try:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(str(color_texto)))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.png'.format(str(self.name)))

    def new_one_color_insert_img_sin_input_with_text_font(self, color, position, inserted, texto, position_texto, color_texto, tu_archivo_de_fuente, tamano_de_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        try:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(str(color_texto)))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.png'.format(str(self.name)))


class JPG:
    def __init__(self, name, width, height):
        self.name = name
        self.width = width
        self.height = height

    def new_one_color(self, color):
        img = Image.new('RGB', (int(self.width), int(self.height)), color=str(color))
        img.save('{}.jpg'.format(str(self.name)))

    def new_one_color_text(self, color, texto, color_texto):
        img = Image.new('RGB', (100, 30), color=(color[0], color[1], color[2]))
        d = ImageDraw.Draw(img)
        d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        img.save('{}.jpg'.format(str(self.name)))

    def new_one_color_text_font(self, color, texto, color_texto, tamano_de_texto, tu_archivo_de_fuente):
        img = Image.new('RGB', (300, 100), color=(color[0], color[1], color[2]))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        d.text((10, 10), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        img.save('{}.jpg'.format(str(self.name)))

    def new_one_color_insert_img(self, color, position):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.jpg'.format(str(self.name)))

    def new_one_color_insert_img_sin_input(self, color, position, inserted):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.jpg'.format(str(self.name)))

    def new_one_color_insert_img_with_text(self, color, position, texto, color_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        d = ImageDraw.Draw(img)
        try:
            d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((10, 10), str(texto), fill=(str(color_texto)))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.jpg'.format(str(self.name)))

    def new_one_color_insert_img_sin_input_with_text(self, color, position, inserted, texto, color_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        d = ImageDraw.Draw(img)
        try:
            d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((10, 10), str(texto), fill=(str(color_texto)))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.jpg'.format(str(self.name)))

    def new_one_color_insert_img_with_text_font(self, color, position, texto, position_texto, color_texto, tu_archivo_de_fuente, tamano_de_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        try:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(str(color_texto)))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.jpg'.format(str(self.name)))

    def new_one_color_insert_img_sin_input_with_text_font(self, color, position, inserted, texto, position_texto, color_texto, tu_archivo_de_fuente, tamano_de_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        try:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(str(color_texto)))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.jpg'.format(str(self.name)))


class ICO:
    def __init__(self, name, width, height):
        self.name = name
        self.width = width
        self.height = height

    def new_one_color(self, color):
        img = Image.new('RGB', (int(self.width), int(self.height)), color=str(color))
        img.save('{}.ico'.format(str(self.name)), format='ICO', sizes=[(int(self.width), int(self.height))])

    def new_one_color_text(self, color, texto, color_texto):
        img = Image.new('RGB', (100, 30), color=(color[0], color[1], color[2]))
        d = ImageDraw.Draw(img)
        d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        img.save('{}.ico'.format(str(self.name)), format='ICO', sizes=[(int(self.width), int(self.height))])

    def new_one_color_text_font(self, color, texto, color_texto, tamano_de_texto, tu_archivo_de_fuente):
        img = Image.new('RGB', (300, 100), color=(color[0], color[1], color[2]))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        d.text((10, 10), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        img.save('{}.ico'.format(str(self.name)), format='ICO', sizes=[(int(self.width), int(self.height))])

    def new_one_color_insert_img(self, color, position):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.ico'.format(str(self.name)), format='ICO', sizes=[(int(self.width), int(self.height))])

    def new_one_color_insert_img_sin_input(self, color, position, inserted):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.ico'.format(str(self.name)), format='ICO', sizes=[(int(self.width), int(self.height))])

    def new_one_color_insert_img_with_text(self, color, position, texto, color_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        d = ImageDraw.Draw(img)
        try:
            d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((10, 10), str(texto), fill=(str(color_texto)))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.ico'.format(str(self.name)), format='ICO', sizes=[(int(self.width), int(self.height))])

    def new_one_color_insert_img_sin_input_with_text(self, color, position, inserted, texto, color_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        d = ImageDraw.Draw(img)
        try:
            d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((10, 10), str(texto), fill=(str(color_texto)))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.ico'.format(str(self.name)), format='ICO', sizes=[(int(self.width), int(self.height))])

    def new_one_color_insert_img_with_text_font(self, color, position, texto, position_texto, color_texto, tu_archivo_de_fuente, tamano_de_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        try:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(str(color_texto)))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.ico'.format(str(self.name)), format='ICO', sizes=[(int(self.width), int(self.height))])

    def new_one_color_insert_img_sin_input_with_text_font(self, color, position, inserted, texto, position_texto, color_texto, tu_archivo_de_fuente, tamano_de_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        try:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(str(color_texto)))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.ico'.format(str(self.name)), format='ICO', sizes=[(int(self.width), int(self.height))])


class PDF:
    def __init__(self, name: str):
        self.nombre = name
        self.documento = Document()
        self.n_cabeceras = 0
        self.n_parrafos = 0
        self.n_imagenes = 0
        self.n_tablas = 0

    def anhadir_cabecera(self, cabecera: str, nivel=1):
        self.documento.add_heading(str(cabecera), level=int(nivel))
        self.n_cabeceras += 1

    def anhadir_parrafo(self, texto: str, estilo_de_letra="Intense Quote"):
        self.documento.add_paragraph(str(texto), style=str(estilo_de_letra))
        self.n_parrafos += 1

    def anhadir_imagen(self, ruta_a_la_imagen: str, ancho=7):
        self.documento.add_picture(str(ruta_a_la_imagen), width=Inches(int(ancho)))
        self.n_imagenes += 1

    def anhadir_lista_ordenada(self, texto: str):
        self.documento.add_paragraph(str(texto), style="List Number")

    def anhadir_lista_desordenada(self, texto: str):
        self.documento.add_paragraph(str(texto), style="List Bullet")

    def anhadir_tabla(self, rows=1, cols=1, cabeceras=(), valores=()):
        table = self.documento.add_table(int(rows), int(cols))
        hdr_cells = table.rows[0].cells
        i = 0
        for cabecera in cabeceras:
            hdr_cells[i].text = str(cabecera)
            i += 1
        for valor in valores:
            row_cells = table.add_row().cells
            i = 0
            while int(i) < int(cols):
                row_cells[int(i)].text = valor[i]
                i += 1
        self.n_tablas += 1

    def guardar(self):
        self.documento.save('{}.docx'.format(str(self.nombre)))
        file = str('{}.docx'.format(str(self.nombre)))
        if str(file) != "":
            convert(file)
        remove("{}.docx".format(str(self.nombre)))

    def guardar_como(self):
        docx_name = asksaveasfilename(defaultextension='.docx', filetypes=[("archivos docx", '*.docx')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        self.documento.save(str(docx_name))
        file = str(docx_name)
        if str(file) != "":
            convert(file)
        remove(str(docx_name))


class DOCX:
    def __init__(self, name: str):
        self.nombre = name
        self.documento = Document()
        self.n_cabeceras = 0
        self.n_parrafos = 0
        self.n_imagenes = 0
        self.n_tablas = 0

    def anhadir_cabecera(self, cabecera: str, nivel=1):
        self.documento.add_heading(str(cabecera), level=int(nivel))
        self.n_cabeceras += 1

    def anhadir_paragrafo(self, texto: str, estilo_de_letra="Intense Quote"):
        self.documento.add_paragraph(str(texto), style=str(estilo_de_letra))
        self.n_parrafos += 1

    def anhadir_imagen(self, ruta_a_la_imagen: str, ancho=7):
        self.documento.add_picture(str(ruta_a_la_imagen), width=Inches(int(ancho)))
        self.n_imagenes += 1

    def anhadir_lista_ordenada(self, texto: str):
        self.documento.add_paragraph(str(texto), style="List Number")

    def anhadir_lista_desordenada(self, texto: str):
        self.documento.add_paragraph(str(texto), style="List Bullet")

    def anhadir_tabla(self, rows=1, cols=1, cabeceras=(), valores=()):
        table = self.documento.add_table(int(rows), int(cols))
        hdr_cells = table.rows[0].cells
        i = 0
        for cabecera in cabeceras:
            hdr_cells[i].text = str(cabecera)
            i += 1
        for valor in valores:
            row_cells = table.add_row().cells
            i = 0
            while int(i) < int(cols):
                row_cells[int(i)].text = valor[i]
                i += 1
        self.n_tablas += 1

    def guardar(self):
        self.documento.save('{}.docx'.format(str(self.nombre)))

    def guardar_como(self):
        docx_name = asksaveasfilename(defaultextension='.docx', filetypes=[("archivos docx", '*.docx')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        self.documento.save(str(docx_name))


class BAT:
    def __init__(self, name, commands):
        self.name = name
        self.commands = commands
        self.n_caracteres = len(commands)

    def new_bat(self):
        bat_file = open("{}.bat".format(str(self.name)), "w")
        bat_file.write(str(self.commands))
        bat_file.close()

    def apend_to_bat(self):
        bat_file = open("{}.bat".format(str(self.name)), "a")
        bat_file.write(str(self.commands))
        bat_file.close()

    def read_bat(self):
        bat_file = open("{}.bat".format(str(self.name)), "r")
        comandos = bat_file.read()
        bat_file.close()
        return comandos


class TXT:
    def __init__(self, name, text):
        self.name = name
        self.text = text
        self.n_caracteres = len(text)

    def new_txt(self):
        txt_file = open("{}.txt".format(str(self.name)), "w")
        txt_file.write(str(self.text))
        txt_file.close()

    def apend_to_txt(self):
        txt_file = open("{}.txt".format(str(self.name)), "a")
        txt_file.write(str(self.text))
        txt_file.close()

    def read_text(self):
        txt_file = open("{}.txt".format(str(self.name)), "r")
        texto = txt_file.read()
        txt_file.close()
        return texto


class WEBP:
    def __init__(self, name, width, height):
        self.name = name
        self.width = width
        self.height = height

    def new_one_color(self, color):
        img = Image.new('RGB', (int(self.width), int(self.height)), color=str(color))
        img.save('{}.webp'.format(str(self.name)), "webp")

    def new_one_color_text(self, color, texto, color_texto):
        img = Image.new('RGB', (100, 30), color=(color[0], color[1], color[2]))
        d = ImageDraw.Draw(img)
        d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        img.save('{}.webp'.format(str(self.name)), "webp")

    def new_one_color_text_font(self, color, texto, color_texto, tamano_de_texto, tu_archivo_de_fuente):
        img = Image.new('RGB', (300, 100), color=(color[0], color[1], color[2]))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        d.text((10, 10), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        img.save('{}.webp'.format(str(self.name)), "webp")

    def new_one_color_insert_img(self, color, position):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.webp'.format(str(self.name)))

    def new_one_color_insert_img_sin_input(self, color, position, inserted):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.webp'.format(str(self.name)))

    def new_one_color_insert_img_with_text(self, color, position, texto, color_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        d = ImageDraw.Draw(img)
        try:
            d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((10, 10), str(texto), fill=(str(color_texto)))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.webp'.format(str(self.name)))

    def new_one_color_insert_img_sin_input_with_text(self, color, position, inserted, texto, color_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        d = ImageDraw.Draw(img)
        try:
            d.text((10, 10), str(texto), fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((10, 10), str(texto), fill=(str(color_texto)))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.webp'.format(str(self.name)))

    def new_one_color_insert_img_with_text_font(self, color, position, texto, position_texto, color_texto, tu_archivo_de_fuente, tamano_de_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        try:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(str(color_texto)))
        im_logo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos de imagen", "*.gif"), ("archivos de imagen", "*.png"), ("archivos de imagen", "*.jpg"), ("archivos de imagen", "*.webp"), ("archivos de imagen", "*.ico")))
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.webp'.format(str(self.name)))

    def new_one_color_insert_img_sin_input_with_text_font(self, color, position, inserted, texto, position_texto, color_texto, tu_archivo_de_fuente, tamano_de_texto):
        try:
            img = Image.new('RGB', (int(self.width), int(self.height)), (color[0], color[1], color[2]))
        except TypeError:
            img = Image.new('RGB', (int(self.width), int(self.height)), str(color))
        fnt = ImageFont.truetype(str(tu_archivo_de_fuente), int(tamano_de_texto))
        d = ImageDraw.Draw(img)
        try:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(color_texto[0], color_texto[1], color_texto[2]))
        except TypeError:
            d.text((int(position_texto[0]), int(position_texto[1])), str(texto), font=fnt, fill=(str(color_texto)))
        im_logo = str(inserted)
        im_logo = Image.open(str(im_logo))
        img.paste(im_logo, (position[0], position[1]))
        img.save('{}.webp'.format(str(self.name)))


class XLSX:
    def __init__(self, ruta):
        self.ruta = ruta

    def new_xlsx_sin_input(self, hojas, datos):
        workbook = xlsxwriter.Workbook('{}.xlsx'.format(self.ruta))
        worksheets = []
        for hoja in hojas:
            worksheets.append(str(hoja))
        for sheet in worksheets:
            worksheet_actual = workbook.add_worksheet(str(sheet))
            for dato in datos:
                if str(sheet) == str(dato[0]):
                    worksheet_actual.write(str(dato[1]), str(dato[2]))
        workbook.close()

    def new_xlsx(self, hojas, datos):
        ruta_al_archivo = asksaveasfilename(defaultextension='.xlsx', filetypes=[("archivos xlsx", '*.xlsx')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        workbook = xlsxwriter.Workbook(str(ruta_al_archivo))
        worksheets = []
        for hoja in hojas:
            worksheets.append(str(hoja))
        for sheet in worksheets:
            worksheet_actual = workbook.add_worksheet(str(sheet))
            for dato in datos:
                if str(sheet) == str(dato[0]):
                    worksheet_actual.write(str(dato[1]), str(dato[2]))
        workbook.close()

    def new_xlsx_grafico(self, hojas, datos, nombre_grafico: str, hoja_grafico: str, categorias_grafico: str, valores_grafico: str, barras: bool, circular: bool, posicion_grafico):
        ruta_al_archivo = asksaveasfilename(defaultextension='.xlsx', filetypes=[("archivos xlsx", '*.xlsx')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        workbook = xlsxwriter.Workbook(str(ruta_al_archivo))
        worksheets = []
        if barras:
            chart = workbook.add_chart({'type': 'column'})
        elif circular:
            chart = workbook.add_chart({'type': 'pie'})
        else:
            chart = workbook.add_chart({'type': 'column'})
        chart.add_series({
            'name': str(nombre_grafico),
            'categories': '=' + str(hoja_grafico) + '!' + str(categorias_grafico),  # '=Sheet1!$A$3:$A$7'
            'values': '=' + str(hoja_grafico) + '!' + str(valores_grafico),  # '=Sheet1!$B$3:$B$7'
            'marker': {'type': 'circle'}
        })
        for hoja in hojas:
            worksheets.append(str(hoja))
        for sheet in worksheets:
            worksheet_actual = workbook.add_worksheet(str(sheet))
            for dato in datos:
                if str(sheet) == str(dato[0]):
                    worksheet_actual.write(str(dato[1]), str(dato[2]))
            if str(sheet) == str(posicion_grafico[0]):
                worksheet_actual.insert_chart(str(posicion_grafico[1]), chart)
        workbook.close()


class CSV:
    def __init__(self, nombre, valores):
        self.nombre = nombre
        self.valores = valores

    def new_csv(self):
        nombre = asksaveasfilename(defaultextension='.csv', filetypes=[("archivos csv", '*.csv')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        csvsalida = open('{}'.format(str(nombre)), 'a', newline='')
        salida = csv.writer(csvsalida)
        salida.writerows(self.valores)
        csvsalida.close()

    def new_csv_sin_input(self):
        csvsalida = open('{}.csv'.format(str(self.nombre)), 'a', newline='')
        salida = csv.writer(csvsalida)
        salida.writerows(self.valores)
        csvsalida.close()

    def read_csv(self):
        nombre = asksaveasfilename(defaultextension='.csv', filetypes=[("archivos csv", '*.csv')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        with open('{}'.format(str(nombre)), newline='') as File:
            reader = csv.reader(File)
            registro = ""
            for row in reader:
                registro = registro + "\n" + str(row)
        print(registro)
        return registro

    def read_csv_sin_input(self):
        with open('{}.csv'.format(str(self.nombre)), newline='') as File:
            reader = csv.reader(File)
            registro = ""
            for row in reader:
                registro = registro + "\n" + str(row)
        print(registro)
        return registro


class JSON:
    def __init__(self, nombre, codigo):
        self.nombre = nombre
        self.codigo = codigo

    def new_json(self):
        nombre = asksaveasfilename(defaultextension='.json', filetypes=[("archivos json", '*.json')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        json_file = open(str(nombre), "a")
        json_file.write(str(self.codigo))
        json_file.close()

    def new_json_sin_input(self):
        json_file = open("{}.json".format(str(self.nombre)), "a")
        json_file.write(str(self.codigo))
        json_file.close()

    def read_json(self):
        json_file = open("{}.json".format(str(self.nombre)), "r")
        content = json_file.read()
        json_file.close()
        return content


class XML:
    def __init__(self, nombre, codigo):
        self.nombre = nombre
        self.codigo = codigo

    def new_xml(self):
        nombre = asksaveasfilename(defaultextension='.xml', filetypes=[("archivos xml", '*.xml')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        xml_file = open(str(nombre), "a")
        xml_file.write(str(self.codigo))
        xml_file.close()

    def new_xml_sin_input(self):
        xml_file = open("{}.xml".format(str(self.nombre)), "a")
        xml_file.write(str(self.codigo))
        xml_file.close()

    def read_xml(self):
        xml_file = open("{}.xml".format(str(self.nombre)), "r")
        content = xml_file.read()
        xml_file.close()
        return content


class JS:
    def __init__(self, name, code):
        self.name = name
        self.code = code

    def new_js(self):
        name = asksaveasfilename(defaultextension='.js', filetypes=[("archivos js", '*.js')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        js_file = open(str(name), "a")
        js_file.write(str(self.code))
        js_file.close()

    def new_js_sin_input(self):
        js_file = open("{}.js".format(str(self.name)), "a")
        js_file.write(str(self.code))
        js_file.close()

    def read_js(self):
        js_file = open("{}.js".format(str(self.name)), "r")
        content = js_file.read()
        js_file.close()
        return content


class PY:
    def __init__(self, name, imports, froms, code, copys):
        self.name = name
        self.imports = imports
        self.froms = froms
        self.code = code
        self.copys = copys

    def new_py(self):
        py_name = asksaveasfilename(defaultextension='.py', filetypes=[("archivos py", '*.py')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        i = 0
        while i < int(self.copys):
            i += 1
            py_file = open("{}.py".format(str(py_name).replace(".py", '') + str(i)), "a")
            codigo = ""
            if self.imports:
                for import1 in self.imports:
                    codigo = codigo + "import " + str(import1) + "\n"
            if self.froms:
                for from1 in self.froms:
                    codigo = codigo + "from " + str(from1[0]) + " " + "import " + str(from1[1]) + "\n"
            codigo = codigo + str(self.code)
            py_file.write(str(codigo))
            py_file.close()

    def new_py_sin_input(self):
        i = 0
        while i < int(self.copys):
            i += 1
            py_file = open("{}.py".format(str(self.name) + str(i)), "a")
            codigo = ""
            if self.imports:
                for import1 in self.imports:
                    codigo = codigo + "import " + str(import1) + "\n"
            if self.froms:
                for from1 in self.froms:
                    codigo = codigo + "from " + str(from1[0]) + " " + "import " + str(from1[1]) + "\n"
            codigo = codigo + str(self.code)
            py_file.write(str(codigo))
            py_file.close()


class CSS:
    def __init__(self, name, styles_body=(), styles_footer=(), styles_p=(), styles_a=(), styles_div=(), styles_span=(), styles_button=()):
        self.styles_body = styles_body
        self.styles_footer = styles_footer
        self.styles_p = styles_p
        self.styles_a = styles_a
        self.styles_div = styles_div
        self.styles_span = styles_span
        self.styles_button = styles_button
        self.name = name

    def new_stylesheet(self):
        css_name = asksaveasfilename(defaultextension='.css', filetypes=[("archivos css", '*.css')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
        css_file = open("{}".format(str(css_name)), "a")
        st_body = "body{\n"
        for style in self.styles_body:
            st_body = st_body + "   {}".format(str(style)) + "\n"
            if style == self.styles_body[len(self.styles_body)-1]:
                st_body = st_body + "}\n\n"
        if st_body != "body{\n":
            css_file.write(st_body)
        del st_body
        st_footer = "footer{\n"
        for style in self.styles_footer:
            st_footer = st_footer + "   {}".format(str(style)) + "\n"
            if style == self.styles_footer[len(self.styles_footer) - 1]:
                st_footer = st_footer + "}\n\n"
        if st_footer != "footer{\n":
            css_file.write(st_footer)
        del st_footer
        st_p = "p{\n"
        for style in self.styles_p:
            st_p = st_p + "   {}".format(str(style)) + "\n"
            if style == self.styles_p[len(self.styles_p) - 1]:
                st_p = st_p + "}\n\n"
        if st_p != "p{\n":
            css_file.write(st_p)
        del st_p
        st_a = "a{\n"
        for style in self.styles_a:
            st_a = st_a + "   {}".format(str(style)) + "\n"
            if style == self.styles_a[len(self.styles_a) - 1]:
                st_a = st_a + "}\n\n"
        if st_a != "a{\n":
            css_file.write(st_a)
        del st_a
        st_div = "div{\n"
        for style in self.styles_div:
            st_div = st_div + "   {}".format(str(style)) + "\n"
            if style == self.styles_div[len(self.styles_div) - 1]:
                st_div = st_div + "}\n\n"
        if st_div != "div{\n":
            css_file.write(st_div)
        del st_div
        st_span = "span{\n"
        for style in self.styles_span:
            st_span = st_span + "   {}".format(str(style)) + "\n"
            if style == self.styles_span[len(self.styles_span) - 1]:
                st_span = st_span + "}\n\n"
        if st_span != "span{\n":
            css_file.write(st_span)
        del st_span
        st_button = "button{\n"
        for style in self.styles_button:
            st_button = st_button + "   {}".format(str(style)) + "\n"
            if style == self.styles_button[len(self.styles_button) - 1]:
                st_button = st_button + "}"
        if st_button != "button{\n":
            css_file.write(st_button)
        del st_button
        css_file.close()

    def new_stylesheet_sin_input(self):
        css_file = open("{}.css".format(str(self.name)), "a")
        st_body = "body{\n"
        for style in self.styles_body:
            st_body = st_body + "   {}".format(str(style)) + "\n"
            if style == self.styles_body[len(self.styles_body)-1]:
                st_body = st_body + "}\n\n"
        if st_body != "body{\n":
            css_file.write(st_body)
        del st_body
        st_footer = "footer{\n"
        for style in self.styles_footer:
            st_footer = st_footer + "   {}".format(str(style)) + "\n"
            if style == self.styles_footer[len(self.styles_footer) - 1]:
                st_footer = st_footer + "}\n\n"
        if st_footer != "footer{\n":
            css_file.write(st_footer)
        del st_footer
        st_p = "p{\n"
        for style in self.styles_p:
            st_p = st_p + "   {}".format(str(style)) + "\n"
            if style == self.styles_p[len(self.styles_p) - 1]:
                st_p = st_p + "}\n\n"
        if st_p != "p{\n":
            css_file.write(st_p)
        del st_p
        st_a = "a{\n"
        for style in self.styles_a:
            st_a = st_a + "   {}".format(str(style)) + "\n"
            if style == self.styles_a[len(self.styles_a) - 1]:
                st_a = st_a + "}\n\n"
        if st_a != "a{\n":
            css_file.write(st_a)
        del st_a
        st_div = "div{\n"
        for style in self.styles_div:
            st_div = st_div + "   {}".format(str(style)) + "\n"
            if style == self.styles_div[len(self.styles_div) - 1]:
                st_div = st_div + "}\n\n"
        if st_div != "div{\n":
            css_file.write(st_div)
        del st_div
        st_span = "span{\n"
        for style in self.styles_span:
            st_span = st_span + "   {}".format(str(style)) + "\n"
            if style == self.styles_span[len(self.styles_span) - 1]:
                st_span = st_span + "}\n\n"
        if st_span != "span{\n":
            css_file.write(st_span)
        del st_span
        st_button = "button{\n"
        for style in self.styles_button:
            st_button = st_button + "   {}".format(str(style)) + "\n"
            if style == self.styles_button[len(self.styles_button) - 1]:
                st_button = st_button + "}"
        if st_button != "button{\n":
            css_file.write(st_button)
        del st_button
        css_file.close()

    def read_css(self):
        css_file = open("{}.css".format(str(self.name)), "r")
        content = css_file.read()
        css_file.close()
        return content


def convertir_xml_json():
    xmlfile = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos xml", "*.xml"), ("archivos xml", "")))
    with open(str(xmlfile)) as xml_file:
        data_dict = xmlparse(xml_file.read())
    xml_file.close()
    json_data = dumps(data_dict)
    jsonfile = asksaveasfilename(defaultextension='.json', filetypes=[("archivos json", '*.json')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
    with open(str(jsonfile), "w") as json_file:
        json_file.write(json_data)
    json_file.close()


def convertir_xml_json_sin_input(ruta_al_archivo, ruta_final_al_json):
    with open(str(ruta_al_archivo)) as xml_file:
        data_dict = xmlparse(xml_file.read())
    xml_file.close()
    json_data = dumps(data_dict)
    with open("{}.json".format(str(ruta_final_al_json)), "w") as json_file:
        json_file.write(json_data)
    json_file.close()


def convertir_json_xml():
    jsonfile = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos json", "*.json"), ("archivos json", "")))
    data = readfromjson(str(jsonfile))
    convertido = json2xml.Json2xml(data).to_xml()
    xmlfile = asksaveasfilename(defaultextension='.xml', filetypes=[("archivos xml", '*.xml')], initialdir="C:/Users/%USERAME%/Documents", title="Choose filename")
    xmlfile = open(str(xmlfile), "a")
    xmlfile.write(convertido)
    xmlfile.close()


def convertir_json_xml_sin_input(ruta_al_archivo, ruta_final_al_xml):
    data = readfromjson(str(ruta_al_archivo))
    convertido = json2xml.Json2xml(data).to_xml()
    xmlfile = open("{}.xml".format(str(ruta_final_al_xml)), "a")
    xmlfile.write(convertido)
    xmlfile.close()


def convertir_png_jpg():
    imagen = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos png", "*.png"), ("archivos png", "")))
    nombre = imagen
    imagen = Image.open(imagen)
    imagen = imagen.convert('RGB')
    imagen.save(str(str(nombre).split("/")[len(str(nombre).split("/"))-1]).replace(".png", "")+".jpg")


def convertir_jpg_png():
    imagen = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos jpg", "*.jpg"), ("archivos jpg", "")))
    nombre = imagen
    imagen = Image.open(imagen)
    imagen.save(str(str(nombre).split("/")[len(str(nombre).split("/"))-1]).replace(".jpg", "")+".png")


def convertir_png_ico():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos png", "*.png"), ("archivos png", "")))
    img = Image.open(filename)
    img.save(str(str(filename).split("/")[len(str(filename).split("/")) - 1]).replace(".png", "") + ".ico", format='ICO', sizes=[(32, 32)])


def convertir_jpg_ico():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos jpg", "*.jpg"), ("archivos jpg", "")))
    img = Image.open(filename)
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".jpg", "")+".ico")


def convertir_ico_png():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos ico", "*.ico"), ("archivos ico", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".ico", "")+".png")


def convertir_ico_jpg():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos ico", "*.ico"), ("archivos ico", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".ico", "")+".jpg")


def convertir_ico_pdf():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos ico", "*.ico"), ("archivos ico", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".ico", "")+".pdf")


def convertir_jpg_pdf():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos jpg", "*.jpg"), ("archivos jpg", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".jpg", "")+".pdf")


def convertir_png_pdf():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos png", "*.png"), ("archivos png", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".png", "")+".pdf")


def convertir_jpg_webp():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos jpg", "*.jpg"), ("archivos jpg", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".jpg", "")+".webp", "webp")


def convertir_png_webp():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos png", "*.png"), ("archivos png", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".png", "")+".webp", "webp")


def convertir_ico_webp():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos ico", "*.ico"), ("archivos ico", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".ico", "")+".webp", "webp")


def convertir_webp_png():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos webp", "*.webp"), ("archivos webp", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".webp", "")+".png")


def convertir_webp_jpg():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos webp", "*.webp"), ("archivos webp", "")))
    img = Image.open(filename)
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".webp", "")+".jpg")


def convertir_webp_ico():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos webp", "*.webp"), ("archivos webp", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/")) - 1]).replace(".webp", "") + ".ico", format='ICO', sizes=[(32, 32)])


def convertir_webp_pdf():
    filename = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos webp", "*.webp"), ("archivos webp", "")))
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".webp", "")+".pdf")


def convertir_txt_xlsx(final_arch):
    arch = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos txt", "*.txt"), ("todos los archivos", "*.*")))
    a = open(str(arch), 'r')
    line = a.read()
    a.close()
    lines = str(line).split()
    workbook = xlsxwriter.Workbook('./{}.xlsx'.format(str(final_arch)))
    worksheet_datos = workbook.add_worksheet('hoja1')
    i = 1
    cad = 0
    cadcont = 1
    num_ac = 1
    for dato in lines:
        letras = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z']
        letra = letras[i - ((26 * cad) + 1)]
        worksheet_datos.write(str(str(letra).upper() + str(num_ac)), dato)
        i += 1
        if cadcont == 26:
            cad += 1
            cadcont = 0
            num_ac += 1
        cadcont += 1
    workbook.close()


def convertir_xlsx_csv(hoja_a_pasar, nombre_csv):
    nombre_xlsx = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos xlsx", "*.xlsx"), ("todos los archivos", "*.*")))
    libro_excel = xlrd.open_workbook('{}'.format(str(nombre_xlsx)))
    hoja = libro_excel.sheet_by_name(str(hoja_a_pasar))
    archivo_csv = open('{}.csv'.format(str(nombre_csv)), 'w')
    escribir_csv = csv.writer(archivo_csv, quoting=csv.QUOTE_ALL)
    for numero in range(hoja.nrows):
        escribir_csv.writerow(hoja.row_values(numero))
    archivo_csv.close()


def convertir_docx_pdf():
    file = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos docx", "*.docx"), ("todos los archivos", "*.*")))
    if str(file) != "":
        convert(file)


def show_img(imagen):
    img = Image.open(str(imagen))
    img.show()


def convertir_docx_txt():
    doc = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos docx", "*.docx"), ("todos los archivos", "*.*")))
    file = docx.Document(doc)
    completo = []
    for paragrafo in file.paragraphs:
        completo.append(paragrafo.text)
    txt = open("pr.txt", "a")
    txt.writelines('\n'.join(completo))
    txt.close()


def convertir_txt_html():
    archivo = askopenfilename(initialdir="C:/Users/%USERNAME%/Documents", title="Selecciona el archivo para convertir", filetypes=(("archivos txt", "*.txt"), ("todos los archivos", "*.*")))
    archivo = open(str(archivo))
    archivo_content = archivo.read()
    archivo.close()
    html = open("mitexto.html", 'a')
    html.write('<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8"><meta name="description" content="El txt que convertí"><title>MyFile.txt</title></head><body><p>{}</p></body></html>'.format(str(archivo_content)))
    html.close()


# A partir de aquí


def convertir_png_jpg_sin_input(ruta_al_archivo):
    imagen = str(ruta_al_archivo)
    nombre = imagen
    imagen = Image.open(imagen)
    imagen = imagen.convert('RGB')
    imagen.save(str(str(nombre).split("/")[len(str(nombre).split("/"))-1]).replace(".png", "")+".jpg")


def convertir_jpg_png_sin_input(ruta_al_archivo):
    imagen = str(ruta_al_archivo)
    nombre = imagen
    imagen = Image.open(imagen)
    imagen.save(str(str(nombre).split("/")[len(str(nombre).split("/"))-1]).replace(".jpg", "")+".png")


def convertir_png_ico_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img.save(str(str(filename).split("/")[len(str(filename).split("/")) - 1]).replace(".png", "") + ".ico", format='ICO', sizes=[(32, 32)])


def convertir_jpg_ico_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".jpg", "")+".ico")


def convertir_ico_png_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".ico", "")+".png")


def convertir_ico_jpg_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".ico", "")+".jpg")


def convertir_ico_pdf_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".ico", "")+".pdf")


def convertir_jpg_pdf_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".jpg", "")+".pdf")


def convertir_png_pdf_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".png", "")+".pdf")


def convertir_jpg_webp_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".jpg", "")+".webp", "webp")


def convertir_png_webp_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".png", "")+".webp", "webp")


def convertir_ico_webp_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".ico", "")+".webp", "webp")


def convertir_webp_png_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".webp", "")+".png")


def convertir_webp_jpg_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".webp", "")+".jpg")


def convertir_webp_ico_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/")) - 1]).replace(".webp", "") + ".ico", format='ICO', sizes=[(32, 32)])


def convertir_webp_pdf_sin_input(ruta_al_archivo):
    filename = str(ruta_al_archivo)
    img = Image.open(filename)
    img = img.convert('RGB')
    img.save(str(str(filename).split("/")[len(str(filename).split("/"))-1]).replace(".webp", "")+".pdf")


def convertir_txt_xlsx_sin_input(ruta_al_archivo, final_arch):
    arch = str(ruta_al_archivo)
    a = open(str(arch), 'r')
    line = a.read()
    a.close()
    lines = str(line).split()
    workbook = xlsxwriter.Workbook('./{}.xlsx'.format(str(final_arch)))
    worksheet_datos = workbook.add_worksheet('hoja1')
    i = 1
    cad = 0
    cadcont = 1
    num_ac = 1
    for dato in lines:
        letras = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z']
        letra = letras[i - ((26 * cad) + 1)]
        worksheet_datos.write(str(str(letra).upper() + str(num_ac)), dato)
        i += 1
        if cadcont == 26:
            cad += 1
            cadcont = 0
            num_ac += 1
        cadcont += 1
    workbook.close()


def convertir_xlsx_csv_sin_input(ruta_al_archivo, hoja_a_pasar, nombre_csv):
    nombre_xlsx = str(ruta_al_archivo)
    libro_excel = xlrd.open_workbook('{}'.format(str(nombre_xlsx)))
    hoja = libro_excel.sheet_by_name(str(hoja_a_pasar))
    archivo_csv = open('{}.csv'.format(str(nombre_csv)), 'w')
    escribir_csv = csv.writer(archivo_csv, quoting=csv.QUOTE_ALL)
    for numero in range(hoja.nrows):
        escribir_csv.writerow(hoja.row_values(numero))
    archivo_csv.close()


def convertir_docx_pdf_sin_input(ruta_al_archivo):
    file = str(ruta_al_archivo)
    if str(file) != "":
        convert(file)


'''
archivo = open("mitexto.txt")
archivo_content = archivo.read()
archivo.close()
print(archivo_content)
'''


def convertir_docx_txt_sin_input(ruta_al_archivo):
    doc = str(ruta_al_archivo)
    file = docx.Document(doc)
    completo = []
    for paragrafo in file.paragraphs:
        completo.append(paragrafo.text)
    txt = open("pr.txt", "a")
    txt.writelines('\n'.join(completo))
    txt.close()


def convertir_txt_html_sin_input(ruta_al_archivo):
    archivo = str(ruta_al_archivo)
    archivo = open(str(archivo))
    archivo_content = archivo.read()
    archivo.close()
    nombre_archivo = str(ruta_al_archivo).split("/")[len(str(ruta_al_archivo).split("/"))-1]
    html = open("mitexto.html", 'a')
    html.write('<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8"><meta name="description" content=""><title>{}</title></head><body><p>{}</p></body></html>'.format(str(nombre_archivo), str(archivo_content)))
    html.close()
